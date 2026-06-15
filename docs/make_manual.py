from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

# 제목
title = doc.add_heading('더원컴퍼니 업무관리시스템', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('사용자 매뉴얼')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(16)
sub.runs[0].bold = True
info = doc.add_paragraph('작성일: 2026-03-23     시스템 주소: http://43.203.193.217:10000')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# 1. 시스템 개요
doc.add_heading('1. 시스템 개요', 1)
doc.add_paragraph('더원컴퍼니 내부 업무 관리 시스템으로, 고객 정보 관리, 공지사항, 인사 관리 등의 기능을 제공합니다.')
t = doc.add_table(rows=4, cols=2)
t.style = 'Table Grid'
data = [('항목','내용'),('시스템 주소','http://43.203.193.217:10000'),
        ('권장 브라우저','Chrome, Edge (최신 버전)'),
        ('문의','더원컴퍼니 | 대전광역시 둔산동 1229 6층 | 02-1660-4130')]
for i,(k,v) in enumerate(data):
    t.cell(i,0).text=k; t.cell(i,1).text=v
    if i==0:
        t.cell(i,0).paragraphs[0].runs[0].bold=True
        t.cell(i,1).paragraphs[0].runs[0].bold=True

# 2. 권한 체계
doc.add_heading('2. 권한 체계', 1)
doc.add_paragraph('사용자는 3가지 권한 중 하나를 가집니다.')
t2 = doc.add_table(rows=4, cols=3)
t2.style = 'Table Grid'
roles = [('권한','명칭','접근 가능 메뉴'),
         ('ADMIN','최고관리자','전체 메뉴 (공지사항 작성, 고객관리, 인사관리, 공통코드관리)'),
         ('MANAGER','매니저','공지사항 작성, 고객관리, 공통코드관리'),
         ('MEMBER','일반사용자','공지사항 읽기, 고객관리')]
for i,row in enumerate(roles):
    for j,v in enumerate(row):
        t2.cell(i,j).text=v
        if i==0: t2.cell(i,j).paragraphs[0].runs[0].bold=True
doc.add_paragraph('※ 권한이 없는 메뉴 접근 시 공지사항 화면으로 이동됩니다.').italic=True

# 3. 로그인/로그아웃
doc.add_heading('3. 로그인 및 로그아웃', 1)
doc.add_heading('3.1 로그인', 2)
for s in ['브라우저에서 시스템 주소로 접속하면 자동으로 로그인 화면이 표시됩니다.',
          '아이디와 비밀번호를 입력 후 로그인 버튼을 클릭합니다.',
          '로그인 성공 시 공지사항 화면으로 이동합니다.']:
    doc.add_paragraph(s, style='List Number')
doc.add_paragraph('오류 상황')
t3 = doc.add_table(rows=4, cols=3)
t3.style = 'Table Grid'
errs = [('상황','메시지','조치'),
        ('아이디/비밀번호 오류','아이디 또는 비밀번호가 올바르지 않습니다','다시 입력'),
        ('비활성 계정','비활성화된 계정입니다','관리자에게 문의'),
        ('비밀번호 초기화됨','비밀번호 재설정 화면으로 자동 이동','4.1 참고')]
for i,row in enumerate(errs):
    for j,v in enumerate(row):
        t3.cell(i,j).text=v
        if i==0: t3.cell(i,j).paragraphs[0].runs[0].bold=True
doc.add_heading('3.2 로그아웃', 2)
doc.add_paragraph('화면 우측 상단 메뉴에서 로그아웃을 클릭합니다. 로그아웃 시 접속 기록이 저장됩니다.')

# 4. 비밀번호 관리
doc.add_heading('4. 비밀번호 관리', 1)
doc.add_heading('4.1 비밀번호 초기화 후 재설정', 2)
doc.add_paragraph('관리자가 비밀번호를 초기화한 경우, 로그인 시 자동으로 비밀번호 재설정 화면이 나타납니다.')
for s in ['새 비밀번호를 입력합니다.','새 비밀번호를 한 번 더 입력합니다.','변경 버튼을 클릭합니다.']:
    doc.add_paragraph(s, style='List Number')
p = doc.add_paragraph()
p.add_run('초기 비밀번호: ').bold=True
p.add_run('아이디 + 1234!  (예: 아이디가 hong이면 hong1234!)')
doc.add_heading('4.2 내 비밀번호 변경', 2)
doc.add_paragraph('비밀번호 규칙:')
for r in ['8자 이상','영문, 숫자, 특수문자를 모두 포함','예시: MyPass1!, Work2026@']:
    doc.add_paragraph(r, style='List Bullet')

# 5. 공지사항
doc.add_heading('5. 공지사항', 1)
doc.add_paragraph('로그인 후 기본으로 표시되는 화면입니다.')
t4 = doc.add_table(rows=4, cols=2)
t4.style = 'Table Grid'
notice = [('기능','설명'),
          ('목록 조회','전체 공지사항이 최신순으로 표시됩니다. 제목 클릭 시 상세 내용 확인 가능'),
          ('작성','ADMIN, MANAGER만 공지사항 작성 가능 (MEMBER는 읽기 전용)'),
          ('수정/삭제','수정: 본인 작성 글만 / 삭제: 본인 작성 글 또는 ADMIN은 전체')]
for i,row in enumerate(notice):
    for j,v in enumerate(row):
        t4.cell(i,j).text=v
        if i==0: t4.cell(i,j).paragraphs[0].runs[0].bold=True

# 6. 고객관리
doc.add_heading('6. 고객관리', 1)
doc.add_heading('6.1 개인고객관리', 2)
doc.add_paragraph('개인 고객 정보를 조회, 등록, 수정, 삭제합니다.')
for s in ['조회: 고객명, 연락처, 상태 등 조건으로 검색 후 조회 버튼 클릭',
          '등록: 등록 버튼 클릭 → 고객 정보 입력 → 저장',
          '수정: 목록에서 고객 선택 → 정보 수정 → 저장',
          '삭제: 목록에서 고객 선택 → 삭제 버튼 클릭']:
    doc.add_paragraph(s, style='List Bullet')
doc.add_heading('6.2 협력점 고객관리', 2)
doc.add_paragraph('협력점(파트너사) 고객 정보를 관리합니다. 조작 방법은 개인고객관리와 동일합니다.')

# 7. 인사관리
doc.add_heading('7. 인사관리 (관리자 전용)', 1)
p = doc.add_paragraph()
p.add_run('ADMIN(최고관리자)만 접근 가능합니다.').bold=True
t5 = doc.add_table(rows=6, cols=2)
t5.style = 'Table Grid'
hr = [('기능','설명'),
      ('직원 목록 조회','아이디, 이름, 부서, 권한, 사용여부로 검색'),
      ('직원 등록','아이디, 비밀번호, 이름, 부서, 권한, 사용여부 입력 후 저장'),
      ('직원 수정','목록 선택 → 정보 수정 → 저장 (비밀번호 입력 시에만 변경)'),
      ('비밀번호 초기화','초기화 버튼 클릭 → 아이디+1234! 로 초기화, 다음 로그인 시 변경 필수'),
      ('직원 삭제','목록 선택 → 삭제 (로그인 이력 포함 삭제)')]
for i,row in enumerate(hr):
    for j,v in enumerate(row):
        t5.cell(i,j).text=v
        if i==0: t5.cell(i,j).paragraphs[0].runs[0].bold=True

# 8. 공통코드관리
doc.add_heading('8. 공통코드관리 (관리자/매니저)', 1)
p = doc.add_paragraph()
p.add_run('ADMIN 및 MANAGER만 접근 가능합니다.').bold=True
doc.add_paragraph('시스템 드롭다운 선택 항목(코드 값)을 관리합니다.')
for s in ['그룹코드 관리: 여러 코드를 묶는 그룹 단위 관리 (조회/등록/수정/삭제)',
          '상세코드 관리: 그룹코드 내 개별 코드 항목 관리',
          '사용여부를 N으로 설정한 코드는 드롭다운 목록에 표시되지 않습니다.']:
    doc.add_paragraph(s, style='List Bullet')

# 9. 세션 관리
doc.add_heading('9. 세션 관리 안내', 1)
t6 = doc.add_table(rows=4, cols=2)
t6.style = 'Table Grid'
sess = [('항목','내용'),
        ('자동 로그아웃','로그인 후 1시간 경과 시 자동 로그아웃'),
        ('만료 안내','"세션이 만료되었습니다. 다시 로그인해주세요." 메시지 표시'),
        ('다중 탭','한 탭에서 로그아웃 시 모든 탭 동시 로그아웃')]
for i,row in enumerate(sess):
    for j,v in enumerate(row):
        t6.cell(i,j).text=v
        if i==0: t6.cell(i,j).paragraphs[0].runs[0].bold=True

# 10. FAQ
doc.add_heading('10. 자주 묻는 질문 (FAQ)', 1)
faqs = [
    ('비밀번호를 잊어버렸습니다.',
     '관리자에게 비밀번호 초기화를 요청하세요. 초기화 후 아이디+1234!로 로그인하면 새 비밀번호로 변경할 수 있습니다.'),
    ('메뉴가 보이지 않습니다.',
     '권한에 따라 접근 가능한 메뉴가 다릅니다. 필요한 권한이 있다면 관리자에게 권한 변경을 요청하세요.'),
    ('로그인이 안 됩니다.',
     '계정이 비활성 상태일 수 있습니다. 관리자에게 계정 활성화를 요청하세요.'),
    ('작업 중 갑자기 로그인 화면으로 이동합니다.',
     '세션(1시간)이 만료된 것입니다. 다시 로그인 후 작업을 계속하세요.'),
    ('공지사항을 작성하려는데 버튼이 없습니다.',
     'MEMBER 권한은 공지사항을 읽기만 할 수 있습니다. 작성은 ADMIN, MANAGER만 가능합니다.'),
]
for q,a in faqs:
    p = doc.add_paragraph()
    p.add_run('Q. ' + q).bold=True
    doc.add_paragraph('A. ' + a)

footer_p = doc.add_paragraph('더원컴퍼니 | 대전광역시 둔산동 1229 6층 | 대표전화: 02-1660-4130')
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = 'C:/Users/ocean-nopc/Desktop/project/managesys/docs/더원컴퍼니_업무관리시스템_사용자매뉴얼.docx'
doc.save(out)
print('워드 저장 완료:', out)
